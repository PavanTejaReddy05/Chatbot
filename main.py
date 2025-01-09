import os
import re
from pymongo import MongoClient
from pathlib import Path
import pdfplumber
import docx
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.responses import JSONResponse, FileResponse
from fastapi.responses import RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi import BackgroundTasks
from pydantic import BaseModel
import cohere
import base64
from gtts import gTTS
from dotenv import load_dotenv
from datetime import datetime, timedelta
from jose import jwt
from passlib.context import CryptContext
from passlib.hash import bcrypt


# client = MongoClient("mongodb+srv://bogireddyptr5:Bogireddyptr5@cluster0.uq0feuh.mongodb.net/")  # Update the URI as needed
# db = client["PDFChatbot"]

# Load environment variables
load_dotenv()

client_str=os.getenv("Client")
client=MongoClient(client_str)

db_name = os.getenv("DB_NAME")
db=client[db_name]

col_name=os.getenv("COLLECTION_NAME")
col = db[col_name]

SECRET_KEY = os.getenv("Secret_Key")
# print(f"Loaded Secret_KEY: {SECRET_KEY}")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 10

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

cohere_api_key = os.getenv("COHERE_API_KEY")
# print(f"Loaded COHERE_API_KEY: {cohere_api_key}")

if not cohere_api_key:
    raise ValueError("COHERE_API_KEY environment variable is not set.")
co = cohere.Client(api_key=cohere_api_key)

app = FastAPI()
# CORS middleware
origins=[
    "*",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_storage():
    return {}

class Question(BaseModel):
    question: str

class Reg(BaseModel):
    First_Name:str
    Last_Name:str
    Email:str
    Ph_No:int
    DoB:str
    Password:str
    Confirm_Password:str

async def Vrfy_Tkn(token:str):
    try:
        payload=jwt.decode(token,SECRET_KEY,algorithms=[ALGORITHM])
        print(f"Token payload: {payload}")
        return payload
    except JWTError as e:
        print(f"Token verification error: {e}")
        return None


@app.post("/Signup")
async def Registration(R: Reg):
    # Check if passwords match
    if R.Password != R.Confirm_Password:
        raise HTTPException(status_code=400, detail="Passwords do not match")

    # Check if the email already exists
    if col.find_one({"Email": R.Email}):
        raise HTTPException(status_code=400, detail="Email already exists")

    # Prepare data for insertion
    data = R.dict()
    data.pop("Confirm_Password")
    data["Password"] = bcrypt.hash(data["Password"])  # Hash the password

    # Insert the user into the database
    try:
        result = col.insert_one(data)
        if result.inserted_id:
            return {"status": "success", "message": "Registration successful"}
        else:
            raise HTTPException(status_code=500, detail="Failed to register user")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.get("/Signin")
def LogIn(ema:str, Pwd:str):
    cred=col.find_one({"Email":ema})
    if not cred:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if bcrypt.verify(Pwd, cred["Password"]):
        pt=datetime.utcnow()
        exp=pt+timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        payload={
            "UserName":ema,
            "present_time":pt.isoformat(),
            "Expiry":exp.isoformat()
            }
        token=jwt.encode(payload,SECRET_KEY,ALGORITHM)
        if token:
            return {"status": "success", "message": "Login successful", "Jwt_Token":token}
    else:
        raise HTTPException(status_code=401, detail="Invalid credentials")

def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    return text

def extract_text_from_docx(file):
    text = ""
    try:
        doc = docx.Document(file)  # Open the DOCX file
        print(f"Number of paragraphs: {len(doc.paragraphs)}")  # Print number of paragraphs for debugging
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)  # Extract and join text from all paragraphs
        if not text:
            print("No text extracted from DOCX file.")
        return text
    except Exception as e:
        print(f"Error while extracting text from DOCX: {str(e)}")
        return text

async def delete_audio_file(file_path: Path):
    """
    Deletes the specified audio file.
    """
    try:
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        print(f"Failed to delete file {file_path}: {e}")

storage = {}

@app.post("/uploadfile/")
async def upload_file(file: UploadFile = File(...)):
    try:
        print(f"Uploaded file: {file.filename}, File extension: {file.filename.split('.')[-1]}")
        file_content = await file.read() 
        print(f"File content size: {len(file_content)} bytes")
        
        # Validate file extension and extract text
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(file.file)
            storage['pdf_text'] = text
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(file.file)
            storage['docx_text'] = text
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        
        # If no text is extracted, raise an error
        if not text:
            raise HTTPException(status_code=400, detail="No readable text found in the document.")

        contact_numbers = re.findall(r'\+?\d[\d -]{8,}\d', text)

        return JSONResponse(content={"contact_numbers": contact_numbers, "full_text": text[:500]})

    except Exception as e:
        print(f"Error occurred: {str(e)}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")


@app.post("/ask")
async def ask_cohere(
    question: Question,
    credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer())
    ):
    global storage  # Access the global storage variable
    token = credentials.credentials
    print(f"Received token: {token}")
    
    payload = await Vrfy_Tkn(token)
    if not payload or datetime.utcnow() >= datetime.fromisoformat(payload.get("Expiry", "")):
        return JSONResponse(
            content={"message": "Token is expired or invalid. Please log in again."},
            status_code=401
        )
        
    text = storage.get('pdf_text') or storage.get('docx_text')  # Check for both PDF and DOCX text
    if not text:
        raise HTTPException(status_code=400, detail="No document content available. Please upload a PDF or DOCX first.")

    try:
        # Prepare the prompt
        prompt = f"""
        Based on the following text from the PDF, answer the question:\n\n
        {text[:2000]}\n\nQuestion: {question.question}
        Answer:
        """
        
        # Log the prompt for debugging
        print(f"Generated prompt: {prompt}")

        # Generate response
        response = co.generate(
            model='command',
            prompt=prompt[:5000],
            max_tokens=300
        )
        
        if not response.generations:
            raise HTTPException(status_code=500, detail="No response from Cohere.")
        
        answer = response.generations[0].text.strip()
        AUDIO_FILE_NAME = "Response.mp3"
        file_path = Path(AUDIO_FILE_NAME)
        if file_path.exists():
            file_path.unlink()  # Delete the file

        tts = gTTS(text=answer, lang='en')
        tts.save(AUDIO_FILE_NAME)
        # Encode audio as base64
        with open(AUDIO_FILE_NAME, "rb") as audio_file:
            encoded_audio = base64.b64encode(audio_file.read()).decode('utf-8')

        # Return JSON response with text and Base64-encoded audio
        return JSONResponse(content={
            "question": question.question,
            "answer": answer,
            "audio_base64": encoded_audio
            })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error occurred: {str(e)}")

